from __future__ import annotations

import math
import os
import sqlite3
from datetime import datetime
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence

import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import sys

_PKG = Path(__file__).resolve().parent.parent
if str(_PKG) not in sys.path:
    sys.path.insert(0, str(_PKG))
from paths import OUTPUT_DIR, ASSET_DIR

MARKET_DB = Path(os.environ.get("MOSS_MARKET_DB_PATH", r"D:\MOSS-SYSTEM-V1\data_warehouse\market.db"))
OUTPUT_DOC = OUTPUT_DIR / f"债券及宏观报告_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
ASSET_DIR.mkdir(exist_ok=True)

COLORS = {
    "navy": "#0B1F33",
    "blue": "#163A63",
    "steel": "#4E6B8A",
    "gold": "#C99A2E",
    "sage": "#6E8B6B",
    "teal": "#2E6F72",
    "orange": "#C76433",
    "rose": "#9E4956",
    "mist": "#EEF3F7",
    "grid": "#C9D4E2",
    "text": "#102235",
    "muted": "#6C7A89",
    "danger": "#B83B3B",
}

NOTE_DEFAULTS = {
    "hs300": (0.10, 0.85),
    "csi500": (0.10, 0.85),
    "gold": (0.15, 0.80),
    "copper": (0.15, 0.80),
    "crude_oil": (0.15, 0.80),
}

ASSET_LABELS = {
    "hs300": "沪深300",
    "csi500": "中证500",
    "gold": "黄金",
    "copper": "铜",
    "crude_oil": "原油",
}


def _set_matplotlib_style() -> None:
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 160
    plt.rcParams["savefig.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"
    plt.rcParams["figure.facecolor"] = "white"


def _styled_axes(ax, title: str, subtitle: Optional[str] = None) -> None:
    ax.set_facecolor("white")
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color(COLORS["grid"])
    ax.spines["bottom"].set_color(COLORS["grid"])
    ax.tick_params(colors=COLORS["muted"], labelsize=9)
    ax.grid(axis="y", color=COLORS["grid"], linestyle="-", linewidth=0.7, alpha=0.7)
    ax.text(0.0, 1.08, title, transform=ax.transAxes, fontsize=13, fontweight="bold", color=COLORS["navy"], ha="left")
    if subtitle:
        ax.text(0.0, 1.02, subtitle, transform=ax.transAxes, fontsize=8.5, color=COLORS["muted"], ha="left")


def _read_sql(query: str) -> pd.DataFrame:
    with sqlite3.connect(MARKET_DB) as conn:
        df = pd.read_sql_query(query, conn)
    if "trade_date" in df.columns:
        df["trade_date"] = pd.to_datetime(df["trade_date"])
    return df


def _load_market_data() -> Dict[str, pd.DataFrame]:
    rates = _read_sql(
        """
        select trade_date, treasury_1y, treasury_3y, treasury_5y, treasury_7y, treasury_10y,
               treasury_30y, term_spread_10y_1y, dr007, reverse_repo_7d
        from market_data_daily
        where trade_date >= date('now', '-420 day')
        order by trade_date
        """
    )
    credit = _read_sql(
        """
        select trade_date, credit_spread_aaa_3y, credit_spread_aa_3y, aa_aaa_spread_3y, cdb_treasury_spread_10y
        from market_data_daily
        where trade_date >= date('now', '-420 day')
        order by trade_date
        """
    )
    macro = _read_sql(
        """
        select trade_date, pmi, cpi_yoy, ppi_yoy, m2_yoy, social_financing_yoy
        from market_data_daily
        where trade_date >= date('now', '-420 day')
        order by trade_date
        """
    )
    return {"rates": rates, "credit": credit, "macro": macro}


def _load_toolkit_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(OUTPUT_DIR / name, encoding="utf-8-sig")


def _latest_valid_row(df: pd.DataFrame, required: Iterable[str]) -> pd.Series:
    mask = df[list(required)].notna().all(axis=1)
    valid = df.loc[mask]
    if valid.empty:
        raise ValueError(f"No valid rows found for {list(required)}")
    return valid.iloc[-1]


def _nearest_prior_row(df: pd.DataFrame, base_date: pd.Timestamp, days: int, required: Iterable[str]) -> pd.Series:
    threshold = base_date - pd.Timedelta(days=days)
    mask = (df["trade_date"] <= threshold) & df[list(required)].notna().all(axis=1)
    valid = df.loc[mask]
    if valid.empty:
        return _latest_valid_row(df[df["trade_date"] <= base_date], required)
    return valid.iloc[-1]


def _monthly_macro(df: pd.DataFrame) -> pd.DataFrame:
    subset = df.dropna(subset=["pmi", "cpi_yoy", "ppi_yoy", "m2_yoy", "social_financing_yoy"], how="all").copy()
    subset["ym"] = subset["trade_date"].dt.to_period("M")
    monthly = subset.sort_values("trade_date").groupby("ym", as_index=False).tail(1)
    return monthly.tail(12).reset_index(drop=True)


def _fmt_num(value: Optional[float], digits: int = 2) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "-"
    return f"{value:.{digits}f}"


def _fmt_change_bp_pct(current: Optional[float], previous: Optional[float]) -> str:
    if current is None or previous is None or any(math.isnan(v) for v in [current, previous]):
        return "-"
    return f'{(current - previous) * 100:+.1f}'


def _fmt_change_bp_level(current: Optional[float], previous: Optional[float]) -> str:
    if current is None or previous is None or any(math.isnan(v) for v in [current, previous]):
        return "-"
    return f'{current - previous:+.1f}'


def _percentile(df: pd.DataFrame, column: str, current: float, days: int = 365) -> Optional[float]:
    recent = df.loc[df["trade_date"] >= df["trade_date"].max() - pd.Timedelta(days=days), column].dropna()
    if recent.empty:
        return None
    return round((recent < current).mean() * 100, 2)


def _safe_float(value) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, str):
        value = value.replace("+", "").strip()
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


@dataclass
class ReportBundle:
    rate_latest: pd.Series
    rate_1m: pd.Series
    rate_3m: pd.Series
    credit_latest: pd.Series
    credit_1m: pd.Series
    credit_3m: pd.Series
    macro_latest: pd.Series
    macro_prev: pd.Series
    rates_df: pd.DataFrame
    credit_df: pd.DataFrame
    macro_monthly: pd.DataFrame
    crisis_latest: pd.Series
    crisis_history: pd.DataFrame
    merrill_latest: pd.Series
    merrill_history: pd.DataFrame
    garch_df: pd.DataFrame
    percentiles: Dict[str, Optional[float]]
    # 新增模型数据
    dcc_latest: Optional[pd.Series]
    risk_parity_df: Optional[pd.DataFrame]
    performance_df: Optional[pd.DataFrame]
    cta_df: Optional[pd.DataFrame]
    rebalance_df: Optional[pd.DataFrame]
    regime_df: Optional[pd.DataFrame]
    backtest_df: Optional[pd.DataFrame]
    backtest_annual_df: Optional[pd.DataFrame]


def _build_bundle() -> ReportBundle:
    market = _load_market_data()
    rates_df = market["rates"]
    credit_df = market["credit"]
    macro_monthly = _monthly_macro(market["macro"])

    rate_latest = _latest_valid_row(rates_df, ["treasury_1y", "treasury_10y", "treasury_30y"])
    rate_1m = _nearest_prior_row(rates_df, rate_latest["trade_date"], 20, ["treasury_1y", "treasury_10y", "treasury_30y"])
    rate_3m = _nearest_prior_row(rates_df, rate_latest["trade_date"], 90, ["treasury_1y", "treasury_10y", "treasury_30y"])

    credit_latest = _latest_valid_row(credit_df, ["credit_spread_aaa_3y", "credit_spread_aa_3y", "aa_aaa_spread_3y"])
    credit_1m = _nearest_prior_row(
        credit_df,
        credit_latest["trade_date"],
        20,
        ["credit_spread_aaa_3y", "credit_spread_aa_3y", "aa_aaa_spread_3y"],
    )
    credit_3m = _nearest_prior_row(
        credit_df,
        credit_latest["trade_date"],
        90,
        ["credit_spread_aaa_3y", "credit_spread_aa_3y", "aa_aaa_spread_3y"],
    )

    crisis_latest = _load_toolkit_csv("crisis_score_latest.csv").iloc[-1]
    crisis_history = _load_toolkit_csv("crisis_score_history.csv")
    crisis_history["date"] = pd.to_datetime(crisis_history["date"])

    merrill_latest = _load_toolkit_csv("merrill_clock_latest.csv").iloc[-1]
    merrill_history = _load_toolkit_csv("merrill_clock_history.csv")
    merrill_history["date"] = pd.to_datetime(merrill_history["date"])

    garch_df = _load_toolkit_csv("garch_results.csv")
    garch_df["asset_name"] = garch_df["资产"].map(lambda x: ASSET_LABELS.get(x, x))
    garch_df["note_alpha"] = garch_df["资产"].map(lambda x: NOTE_DEFAULTS[x][0])
    garch_df["note_beta"] = garch_df["资产"].map(lambda x: NOTE_DEFAULTS[x][1])

    percentiles = {
        "y10_1y": _percentile(rates_df, "treasury_10y", float(rate_latest["treasury_10y"])),
        "curve_1y": _percentile(rates_df, "term_spread_10y_1y", float(rate_latest["term_spread_10y_1y"])),
        "aaa_1y": _percentile(credit_df, "credit_spread_aaa_3y", float(credit_latest["credit_spread_aaa_3y"])),
        "aa_aaa_1y": _percentile(credit_df, "aa_aaa_spread_3y", float(credit_latest["aa_aaa_spread_3y"])),
    }

    # 新模型 CSV（可选，文件不存在时为 None）
    def _try_load(name: str) -> Optional[pd.DataFrame]:
        p = OUTPUT_DIR / name
        return pd.read_csv(p, encoding="utf-8-sig") if p.exists() else None

    dcc_df = _try_load("dcc_latest.csv")
    dcc_latest = dcc_df.iloc[-1] if dcc_df is not None and len(dcc_df) else None

    risk_parity_df = _try_load("risk_parity_results.csv")
    performance_df = _try_load("performance_results.csv")
    cta_df         = _try_load("cta_results.csv")
    rebalance_df   = _try_load("rebalance_results.csv")
    regime_df      = _try_load("regime_results.csv")
    backtest_df    = _try_load("backtest_results.csv")
    backtest_annual_df = _try_load("backtest_annual.csv")

    return ReportBundle(
        rate_latest=rate_latest,
        rate_1m=rate_1m,
        rate_3m=rate_3m,
        credit_latest=credit_latest,
        credit_1m=credit_1m,
        credit_3m=credit_3m,
        macro_latest=macro_monthly.iloc[-1],
        macro_prev=macro_monthly.iloc[-2],
        rates_df=rates_df,
        credit_df=credit_df,
        macro_monthly=macro_monthly,
        crisis_latest=crisis_latest,
        crisis_history=crisis_history,
        merrill_latest=merrill_latest,
        merrill_history=merrill_history,
        garch_df=garch_df,
        percentiles=percentiles,
        dcc_latest=dcc_latest,
        risk_parity_df=risk_parity_df,
        performance_df=performance_df,
        cta_df=cta_df,
        rebalance_df=rebalance_df,
        regime_df=regime_df,
        backtest_df=backtest_df,
        backtest_annual_df=backtest_annual_df,
    )


def _plot_yield_curve(bundle: ReportBundle) -> Path:
    path = ASSET_DIR / "yield_curve.png"
    tenors = ["1Y", "3Y", "5Y", "7Y", "10Y", "30Y"]
    cols = {"1Y": "treasury_1y", "3Y": "treasury_3y", "5Y": "treasury_5y", "7Y": "treasury_7y", "10Y": "treasury_10y", "30Y": "treasury_30y"}
    fig, ax = plt.subplots(figsize=(8.5, 4.9))
    _styled_axes(ax, "国债收益率曲线", "对比当前、1个月前和3个月前的期限结构变化")
    rows = [
        (bundle.rate_3m, bundle.rate_3m["trade_date"].strftime("%Y-%m-%d"), COLORS["steel"], 1.8, 0.75),
        (bundle.rate_1m, bundle.rate_1m["trade_date"].strftime("%Y-%m-%d"), COLORS["gold"], 2.0, 0.9),
        (bundle.rate_latest, bundle.rate_latest["trade_date"].strftime("%Y-%m-%d"), COLORS["navy"], 2.8, 1.0),
    ]
    for row, label, color, width, alpha in rows:
        values = [float(row[cols[t]]) for t in tenors]
        ax.plot(tenors, values, marker="o", linewidth=width, markersize=5.5, color=color, alpha=alpha, label=label)
    latest_vals = [float(bundle.rate_latest[cols[t]]) for t in tenors]
    ax.fill_between(tenors, latest_vals, [min(latest_vals) - 0.1] * len(tenors), color=COLORS["mist"], alpha=0.8)
    ax.set_ylabel("收益率（%）", color=COLORS["text"])
    ax.legend(loc="upper left", frameon=False, fontsize=8.5)
    ax.annotate(
        f"10Y: {float(bundle.rate_latest['treasury_10y']):.4f}%",
        xy=("10Y", float(bundle.rate_latest["treasury_10y"])),
        xytext=(12, 12),
        textcoords="offset points",
        fontsize=8.5,
        color=COLORS["navy"],
        bbox=dict(boxstyle="round,pad=0.25", fc="white", ec=COLORS["grid"]),
    )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def _plot_liquidity(bundle: ReportBundle) -> Path:
    path = ASSET_DIR / "liquidity.png"
    df = bundle.rates_df.dropna(subset=["dr007", "reverse_repo_7d"], how="all").tail(140).copy()
    df["spread_bp"] = (df["dr007"] - df["reverse_repo_7d"]) * 100
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(8.5, 6.0), gridspec_kw={"height_ratios": [2.2, 1]}, sharex=True)
    _styled_axes(ax1, "资金面与政策利率", "DR007 相对 7 天逆回购利率的偏离决定了利率债的短期舒适度")
    _styled_axes(ax2, "流动性偏离", None)
    ax1.plot(df["trade_date"], df["dr007"], color=COLORS["navy"], linewidth=2.4, label="DR007")
    ax1.plot(df["trade_date"], df["reverse_repo_7d"], color=COLORS["gold"], linewidth=2.0, label="7天逆回购")
    ax1.fill_between(df["trade_date"], df["dr007"], df["reverse_repo_7d"], color=COLORS["mist"], alpha=0.8)
    ax1.legend(loc="upper left", frameon=False, fontsize=8.5)
    ax1.set_ylabel("利率（%）")
    ax2.bar(df["trade_date"], df["spread_bp"], width=1.6, color=np.where(df["spread_bp"] >= 0, COLORS["orange"], COLORS["sage"]), alpha=0.9)
    ax2.axhline(0, color=COLORS["grid"], linewidth=1)
    ax2.set_ylabel("bp")
    ax2.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
    fig.autofmt_xdate(rotation=30)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def _plot_credit(bundle: ReportBundle) -> Path:
    path = ASSET_DIR / "credit_spreads.png"
    df = bundle.credit_df.dropna(subset=["credit_spread_aaa_3y", "aa_aaa_spread_3y"], how="all").tail(180).copy()
    fig, ax = plt.subplots(figsize=(8.5, 4.9))
    _styled_axes(ax, "信用利差与等级分化", "AAA 利差反映整体风险偏好，AA-AAA 反映下沉信用的承受能力")
    ax.plot(df["trade_date"], df["credit_spread_aaa_3y"], color=COLORS["navy"], linewidth=2.4, label="AAA 3Y信用利差")
    ax.plot(df["trade_date"], df["aa_aaa_spread_3y"], color=COLORS["rose"], linewidth=2.2, label="AA-AAA等级利差")
    ax.fill_between(df["trade_date"], df["credit_spread_aaa_3y"], color=COLORS["mist"], alpha=0.4)
    ax.set_ylabel("BP")
    ax.legend(loc="upper left", frameon=False, fontsize=8.5)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
    fig.autofmt_xdate(rotation=30)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def _plot_macro(bundle: ReportBundle) -> Path:
    path = ASSET_DIR / "macro_monthly.png"
    df = bundle.macro_monthly.copy()
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(8.5, 6.3), gridspec_kw={"height_ratios": [1.2, 1.8]}, sharex=True)
    _styled_axes(ax1, "宏观动量面板", "PMI 与价格、信用总量并置观察，更接近交易桌跟踪逻辑")
    _styled_axes(ax2, "", None)
    ax1.plot(df["trade_date"], df["pmi"], color=COLORS["navy"], marker="o", linewidth=2.4, label="PMI")
    ax1.axhline(50, color=COLORS["gold"], linestyle="--", linewidth=1.2)
    ax1.fill_between(df["trade_date"], df["pmi"], 50, color=COLORS["mist"], alpha=0.55)
    ax1.set_ylabel("PMI")
    ax1.legend(loc="upper left", frameon=False, fontsize=8.5)
    ax2.plot(df["trade_date"], df["cpi_yoy"], color=COLORS["orange"], marker="o", linewidth=2.0, label="CPI同比")
    ax2.plot(df["trade_date"], df["ppi_yoy"], color=COLORS["rose"], marker="o", linewidth=2.0, label="PPI同比")
    ax2.plot(df["trade_date"], df["m2_yoy"], color=COLORS["teal"], marker="o", linewidth=2.0, label="M2同比")
    ax2.plot(df["trade_date"], df["social_financing_yoy"], color=COLORS["sage"], marker="o", linewidth=2.0, label="社融同比")
    ax2.set_ylabel("%")
    ax2.legend(loc="upper left", frameon=False, fontsize=8.5, ncol=2)
    ax2.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
    fig.autofmt_xdate(rotation=30)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def _plot_garch_params(bundle: ReportBundle) -> Path:
    path = ASSET_DIR / "garch_params.png"
    df = bundle.garch_df.copy()
    x = np.arange(len(df))
    width = 0.18
    fig, axes = plt.subplots(2, 1, figsize=(8.5, 6.5), sharex=True)
    _styled_axes(axes[0], "GARCH 参数对照", "实际估计值与笔记默认值对比：alpha 决定对新信息反应，beta 决定波动持续性")
    _styled_axes(axes[1], "波动持续性差异", None)
    axes[0].bar(x - width / 2, df["alpha"], width=width, color=COLORS["navy"], label="实际 alpha")
    axes[0].bar(x + width / 2, df["note_alpha"], width=width, color=COLORS["gold"], label="默认 alpha")
    axes[0].bar(x + width * 1.8, df["beta"], width=width, color=COLORS["rose"], label="实际 beta")
    axes[0].bar(x + width * 2.8, df["note_beta"], width=width, color=COLORS["sage"], label="默认 beta")
    axes[0].legend(loc="upper left", frameon=False, fontsize=8.5, ncol=2)
    axes[0].set_ylabel("参数值")
    axes[0].set_ylim(0, max(df["beta"].max(), df["note_beta"].max()) + 0.18)
    beta_diff = df["beta"] - df["note_beta"]
    axes[1].bar(x, beta_diff, color=[COLORS["danger"] if v > 0 else COLORS["teal"] for v in beta_diff], width=0.5)
    axes[1].axhline(0, color=COLORS["grid"], linewidth=1)
    axes[1].set_ylabel("实际β-默认β")
    axes[1].set_xticks(x + width)
    axes[1].set_xticklabels(df["asset_name"], fontsize=9, color=COLORS["text"])
    for idx, val in enumerate(beta_diff):
        axes[1].text(idx, val + (0.015 if val >= 0 else -0.05), f'{val:+.3f}', ha="center", fontsize=8.5, color=COLORS["text"])
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def _plot_garch_risk(bundle: ReportBundle) -> Path:
    path = ASSET_DIR / "garch_risk.png"
    df = bundle.garch_df.sort_values("年化波动率%", ascending=True).copy()
    fig, ax = plt.subplots(figsize=(8.5, 4.9))
    _styled_axes(ax, "跨资产波动状态", "年化波动率越高，越不适合在当前阶段承担方向性进攻仓位")
    palette = []
    for state in df["波动率状态"]:
        if "极端" in state:
            palette.append(COLORS["danger"])
        elif "高" in state:
            palette.append(COLORS["orange"])
        else:
            palette.append(COLORS["navy"])
    ax.barh(df["asset_name"], df["年化波动率%"], color=palette, alpha=0.92, height=0.55)
    ax.set_xlabel("年化波动率（%）")
    ax.set_ylabel("")
    for _, row in df.iterrows():
        ax.text(
            float(row["年化波动率%"]) + 1.0,
            row["asset_name"],
            f"{float(row['年化波动率%']):.2f}% | {row['波动率状态']}",
            va="center",
            fontsize=8.5,
            color=COLORS["text"],
        )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_document_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(11, 31, 51)


def _add_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Microsoft YaHei"
    run.font.color.rgb = RGBColor(11, 31, 51)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(10.5)
    run2.font.name = "Microsoft YaHei"
    run2.font.color.rgb = RGBColor(108, 122, 137)
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_bullets(document: Document, items: Sequence[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _add_table(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        hdr[idx].text = str(col)
        _set_cell_shading(hdr[idx], "DCE6F1")
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def _macro_takeaways(bundle: ReportBundle) -> List[str]:
    pmi = float(bundle.macro_latest["pmi"])
    cpi = float(bundle.macro_latest["cpi_yoy"])
    ppi = float(bundle.macro_latest["ppi_yoy"])
    m2 = float(bundle.macro_latest["m2_yoy"])
    sf = float(bundle.macro_latest["social_financing_yoy"])
    merrill_phase = str(bundle.merrill_latest["传统象限"])
    crisis_score = _safe_float(bundle.crisis_latest.get("Crisis Score")) or 0.0

    # 增长描述
    if pmi >= 51:
        growth_desc = "明显扩张"
    elif pmi >= 50:
        growth_desc = "温和扩张"
    elif pmi >= 49:
        growth_desc = "轻微收缩"
    else:
        growth_desc = "明显收缩"

    # 价格描述
    if cpi < 0 and ppi < 0:
        price_desc = "CPI 和 PPI 双双为负，通缩压力较重"
    elif cpi < 1 and ppi < 0:
        price_desc = "CPI 低位、PPI 仍负，通缩压力尚未消散"
    elif cpi >= 2 or ppi >= 3:
        price_desc = "价格动量明显抬升，通胀压力值得关注"
    else:
        price_desc = "价格温和，通胀压力可控"

    # M2/社融描述
    if m2 >= 10:
        liquidity_desc = "货币总量宽松，流动性充裕"
    elif m2 >= 8:
        liquidity_desc = "货币总量适中，流动性平稳"
    else:
        liquidity_desc = "货币总量偏紧，流动性边际收敛"

    # 美林时钟象限解读
    phase_desc = {
        "复苏": "增长动量回升、通胀动量仍低，利好股票和信用债",
        "过热": "增长和通胀动量双升，利好商品，债券承压",
        "滞胀": "增长动量走弱、通胀动量抬升，利好黄金和现金，债券和股票均承压",
        "衰退": "增长和通胀动量双降，利好债券，股票和商品承压",
    }.get(merrill_phase, f"当前处于{merrill_phase}象限")

    # Crisis Score 描述
    if crisis_score < 0:
        crisis_desc = f"Crisis Score 为 {crisis_score:.3f}，市场处于宽松区间，系统性风险极低"
    elif crisis_score < 1:
        crisis_desc = f"Crisis Score 为 {crisis_score:.3f}，市场处于正常区间，无明显系统性压力"
    elif crisis_score < 2:
        crisis_desc = f"Crisis Score 为 {crisis_score:.3f}，市场进入警惕区间，需关注风险积累"
    elif crisis_score < 3:
        crisis_desc = f"Crisis Score 为 {crisis_score:.3f}，市场处于高风险区间，建议降低风险敞口"
    else:
        crisis_desc = f"Crisis Score 为 {crisis_score:.3f}，市场处于危机区间，建议大幅防御"

    return [
        f"最新月度宏观读数显示增长端{growth_desc}，PMI 为 {pmi:.1f}；{price_desc}，CPI 同比 {cpi:.1f}%，PPI 同比 {ppi:.1f}%。",
        f"{liquidity_desc}，M2 同比 {m2:.1f}%，社融存量同比 {sf:.1f}%。",
        f"中国版美林时钟当前落在'{merrill_phase}'：{phase_desc}。",
        f"{crisis_desc}。",
    ]


def _bond_takeaways(bundle: ReportBundle) -> List[str]:
    y10 = float(bundle.rate_latest["treasury_10y"])
    y10_1m = float(bundle.rate_1m["treasury_10y"])
    curve = float(bundle.rate_latest["term_spread_10y_1y"])
    dr007 = float(bundle.rate_latest["dr007"])
    repo = float(bundle.rate_latest["reverse_repo_7d"])
    y10_pct = bundle.percentiles["y10_1y"] or 50.0
    curve_pct = bundle.percentiles["curve_1y"] or 50.0

    # 10Y 变动方向
    y10_chg_bp = (y10 - y10_1m) * 100
    y10_chg_desc = f"上行 {y10_chg_bp:.1f}bp" if y10_chg_bp > 0 else f"下行 {abs(y10_chg_bp):.1f}bp"

    # 收益率分位数解读
    if y10_pct < 20:
        y10_pct_desc = f"位于近1年极低分位（{y10_pct:.0f}%），利率已充分定价，追多性价比低"
        trade_view = "当前利率处于历史低位，建议以持有为主，不宜激进追多"
    elif y10_pct < 40:
        y10_pct_desc = f"位于近1年偏低分位（{y10_pct:.0f}%），仍有一定配置价值"
        trade_view = "利率偏低但未到极端，可逢调整适度加仓"
    elif y10_pct < 60:
        y10_pct_desc = f"位于近1年中性分位（{y10_pct:.0f}%），利率处于合理区间"
        trade_view = "利率处于中性区间，持有为主，等待方向明确"
    elif y10_pct < 80:
        y10_pct_desc = f"位于近1年偏高分位（{y10_pct:.0f}%），利率有一定配置吸引力"
        trade_view = "利率偏高，具备配置价值，可适度增加久期"
    else:
        y10_pct_desc = f"位于近1年高分位（{y10_pct:.0f}%），利率处于历史高位，配置价值突出"
        trade_view = "利率处于历史高位，建议积极配置，拉长久期"

    # 曲线形态
    if curve_pct > 60:
        curve_desc = f"曲线偏陡（{curve:.2f}bp，近1年{curve_pct:.0f}%分位），长端相对短端仍有溢价"
    elif curve_pct < 30:
        curve_desc = f"曲线偏平（{curve:.2f}bp，近1年{curve_pct:.0f}%分位），长短端利差压缩，期限溢价不足"
    else:
        curve_desc = f"曲线形态中性（{curve:.2f}bp，近1年{curve_pct:.0f}%分位）"

    # 资金面
    spread_bp = (dr007 - repo) * 100
    if dr007 < repo:
        liquidity_desc = f"DR007（{dr007:.4f}%）低于逆回购（{repo:.4f}%），资金面偏松，短端有支撑"
    elif spread_bp < 10:
        liquidity_desc = f"DR007（{dr007:.4f}%）略高于逆回购（{repo:.4f}%），资金面基本平衡"
    else:
        liquidity_desc = f"DR007（{dr007:.4f}%）明显高于逆回购（{repo:.4f}%），资金面偏紧，需关注短端压力"

    return [
        f"截至 {bundle.rate_latest['trade_date'].strftime('%Y-%m-%d')}，10 年国债收益率为 {y10:.4f}%，较约 1 个月前{y10_chg_desc}，{y10_pct_desc}。",
        f"{curve_desc}。",
        f"{liquidity_desc}。",
        f"{trade_view}。",
    ]


def _credit_takeaways(bundle: ReportBundle) -> List[str]:
    aaa = float(bundle.credit_latest["credit_spread_aaa_3y"])
    aa_aaa = float(bundle.credit_latest["aa_aaa_spread_3y"])
    aaa_pct = bundle.percentiles["aaa_1y"] or 50.0
    aa_aaa_pct = bundle.percentiles["aa_aaa_1y"] or 50.0
    crisis = bundle.crisis_latest
    fx_z     = _safe_float(crisis.get("汇率波动率z")) or 0.0
    commod_z = _safe_float(crisis.get("商品波动率z")) or 0.0
    credit_z = _safe_float(crisis.get("信用利差z"))  or 0.0

    # AAA 利差分位数解读
    if aaa_pct < 20:
        aaa_desc = f"AAA 3Y 信用利差 {aaa:.2f}bp，位于近1年极低分位（{aaa_pct:.0f}%），高等级信用已明显压缩，性价比偏低"
    elif aaa_pct < 40:
        aaa_desc = f"AAA 3Y 信用利差 {aaa:.2f}bp，位于近1年偏低分位（{aaa_pct:.0f}%），高等级信用利差偏窄"
    elif aaa_pct < 70:
        aaa_desc = f"AAA 3Y 信用利差 {aaa:.2f}bp，位于近1年中性分位（{aaa_pct:.0f}%），高等级信用估值合理"
    else:
        aaa_desc = f"AAA 3Y 信用利差 {aaa:.2f}bp，位于近1年偏高分位（{aaa_pct:.0f}%），高等级信用具备配置价值"

    # AA-AAA 等级利差解读
    if aa_aaa_pct > 70:
        grade_desc = f"AA-AAA 等级利差 {aa_aaa:.2f}bp（近1年{aa_aaa_pct:.0f}%分位），下沉信用风险补偿偏高，市场对弱资质主体仍保持谨慎"
        credit_strategy = "建议以高等级为主，控制下沉节奏，等待等级利差收窄信号"
    elif aa_aaa_pct < 30:
        grade_desc = f"AA-AAA 等级利差 {aa_aaa:.2f}bp（近1年{aa_aaa_pct:.0f}%分位），等级利差已明显压缩，下沉信用性价比下降"
        credit_strategy = "等级利差已压缩，不建议继续下沉，维持高等级配置"
    else:
        grade_desc = f"AA-AAA 等级利差 {aa_aaa:.2f}bp（近1年{aa_aaa_pct:.0f}%分位），等级分化处于中性水平"
        credit_strategy = "信用分层合理，可适度参与高等级城投和央国企产业债"

    # 外部风险分项
    risk_items = []
    if abs(fx_z) > 1.5:
        risk_items.append(f"汇率波动率 z 值 {fx_z:.2f}（偏高）")
    if abs(commod_z) > 1.5:
        risk_items.append(f"商品波动率 z 值 {commod_z:.2f}（偏高）")
    if abs(credit_z) > 1.5:
        risk_items.append(f"信用利差 z 值 {credit_z:.2f}（偏高）")
    risk_desc = "Crisis Score 分项中，" + "、".join(risk_items) + "，需关注外部冲击传导" if risk_items else "Crisis Score 各分项均处于正常区间，无明显外部冲击压力"

    return [
        f"{aaa_desc}；{grade_desc}。",
        f"{credit_strategy}。",
        f"{risk_desc}。",
    ]


def _garch_takeaways(bundle: ReportBundle) -> List[str]:
    extreme = bundle.garch_df[bundle.garch_df["波动率状态"].astype(str).str.contains("极端")]["asset_name"].tolist()
    extreme_text = "、".join(extreme) if extreme else "无"
    return [
        "GARCH 参数不能直接套笔记默认值。沪深300的 alpha 低于默认股票参数，说明对新信息的反应比预设更迟钝；如果沿用默认值，会高估市场对突发冲击的即时反应。",
        "黄金的 beta 明显高于默认商品参数，同时模型自动选择了 EGARCH / t，说明它不仅波动持续性更强，而且存在显著的不对称效应。",
        "中证500与原油选择了 skew-t，意味着收益率分布偏斜、左尾更厚；铜最接近教科书式商品，参数与默认值最接近。",
        f"当前亮红灯的资产是 {extreme_text}。其中黄金年化波动率约 {bundle.garch_df.loc[bundle.garch_df['asset_name']=='黄金','年化波动率%'].iloc[0]:.2f}%，原油约 {bundle.garch_df.loc[bundle.garch_df['asset_name']=='原油','年化波动率%'].iloc[0]:.2f}%，都不适合承担过重方向性仓位。",
        f"样本外相关性最低的资产仍有 {bundle.garch_df['样本外相关性'].min():.3f}，说明这套波动率模型整体是可用的，足以作为风控和仓位约束信号。",
    ]


def _model_cross_takeaways(bundle: ReportBundle) -> List[str]:
    crisis_score = _safe_float(bundle.crisis_latest.get("Crisis Score")) or 0.0
    fx_z     = _safe_float(bundle.crisis_latest.get("汇率波动率z"))  or 0.0
    commod_z = _safe_float(bundle.crisis_latest.get("商品波动率z"))  or 0.0
    equity_z = _safe_float(bundle.crisis_latest.get("股市波动率z"))  or 0.0
    merrill_phase = str(bundle.merrill_latest["传统象限"])
    bond_pref = _safe_float(bundle.merrill_latest["债券偏好"]) or 0.0
    stock_pref = _safe_float(bundle.merrill_latest["股票偏好"]) or 0.0
    gold_pref = _safe_float(bundle.merrill_latest["黄金偏好"]) or 0.0

    # GARCH 信号
    extreme = bundle.garch_df[bundle.garch_df["波动率状态"].astype(str).str.contains("极端")]["asset_name"].tolist()
    if extreme:
        garch_signal = f"GARCH 显示 {'、'.join(extreme)} 处于极端波动状态，方向性仓位需谨慎"
    else:
        garch_signal = "GARCH 各资产波动率均处于正常区间，无极端波动预警"

    # Crisis Score 信号
    if crisis_score < 1:
        crisis_signal = f"Crisis Score {crisis_score:.3f}（正常），系统性风险可控"
    elif crisis_score < 2:
        crisis_signal = f"Crisis Score {crisis_score:.3f}（警惕），风险在积累"
    else:
        crisis_signal = f"Crisis Score {crisis_score:.3f}（高风险），建议防御"

    # 主要压力来源
    pressure = []
    if abs(fx_z) > 1.5:   pressure.append(f"汇率（z={fx_z:.2f}）")
    if abs(commod_z) > 1.5: pressure.append(f"商品（z={commod_z:.2f}）")
    if abs(equity_z) > 1.5: pressure.append(f"股市（z={equity_z:.2f}）")
    pressure_desc = "压力主要来自 " + "、".join(pressure) if pressure else "各分项均无明显压力"

    # 美林时钟资产排序（取偏好分最高的前3）
    prefs = [
        ("黄金", gold_pref),
        ("股票", stock_pref),
        ("债券", bond_pref),
        ("现金", _safe_float(bundle.merrill_latest["现金偏好"]) or 0.0),
        ("商品", _safe_float(bundle.merrill_latest["商品偏好"]) or 0.0),
    ]
    prefs_sorted = sorted(prefs, key=lambda x: x[1], reverse=True)
    top3 = " > ".join(f"{n}({v:+.2f})" for n, v in prefs_sorted[:3])

    # 三模型综合判断
    defensive_signals = sum([
        crisis_score > 1,
        bool(extreme),
        bond_pref < -0.2,
        stock_pref < -0.2,
    ])
    if defensive_signals >= 3:
        overall = "三模型信号高度一致偏防御，当前不是进攻窗口，建议降低风险敞口"
    elif defensive_signals >= 2:
        overall = "三模型信号偏防御，建议维持中性偏保守仓位，等待风险信号缓解"
    elif defensive_signals == 1:
        overall = "三模型信号分歧，部分防御信号出现，建议维持中性仓位"
    else:
        overall = "三模型信号偏积极，风险环境相对友好，可适度增加风险敞口"

    return [
        f"{garch_signal}；{crisis_signal}，{pressure_desc}。",
        f"美林时钟处在'{merrill_phase}'，资产偏好排序：{top3}。",
        f"{overall}。",
    ]


def _quant_models_takeaways(bundle: ReportBundle) -> List[str]:
    """六个量化模型的综合解读"""
    lines = []

    # DCC
    if bundle.dcc_latest is not None:
        avg_corr = _safe_float(bundle.dcc_latest.get("平均相关系数"))
        warning  = str(bundle.dcc_latest.get("预警状态", "正常"))
        if avg_corr is not None:
            lines.append(
                f"DCC 动态相关：当前多资产平均相关系数 {avg_corr:.3f}，状态为'{warning}'。"
                + ("相关性偏高，分散化效果减弱，需警惕组合集中风险。" if "预警" in warning else "相关性处于合理水平，分散化仍有效。")
            )

    # 风险平价
    if bundle.risk_parity_df is not None and len(bundle.risk_parity_df):
        top_asset = bundle.risk_parity_df.sort_values("风险平价权重%", ascending=False).iloc[0]
        lines.append(
            f"风险平价：当前最大权重资产为'{top_asset['资产']}'（{float(top_asset['风险平价权重%']):.1f}%），"
            f"风险预算权重受美林时钟象限动态调整。"
        )

    # 绩效
    if bundle.performance_df is not None and len(bundle.performance_df):
        perf_col = "资产/组合" if "资产/组合" in bundle.performance_df.columns else "资产"
        # 只取单资产行（排除"组合"行）
        asset_perf = bundle.performance_df[~bundle.performance_df[perf_col].str.contains("组合", na=False)]
        if len(asset_perf):
            best = asset_perf.sort_values("夏普比率", ascending=False).iloc[0]
            lines.append(
                f"绩效评估：近3年夏普比率最高的资产为'{best[perf_col]}'（{float(best['夏普比率']):.3f}），"
                f"年化收益 {float(best['年化收益%']):.2f}%，年化波动 {float(best['年化波动%']):.2f}%。"
            )

    # CTA
    if bundle.cta_df is not None and len(bundle.cta_df):
        strongest = bundle.cta_df.sort_values("合成信号", key=abs, ascending=False).iloc[0]
        lines.append(
            f"CTA 趋势：当前趋势最强资产为'{strongest['资产']}'（合成信号 {float(strongest['合成信号']):+.3f}，{strongest['操作建议']}），"
            f"策略近2年夏普比率 {strongest['策略夏普比率']}。"
        )

    # 再平衡
    if bundle.rebalance_df is not None and len(bundle.rebalance_df):
        best_rebal = bundle.rebalance_df.sort_values("夏普比率", ascending=False).iloc[0]
        lines.append(
            f"再平衡策略：最优方案为'{best_rebal['策略']}'（夏普 {float(best_rebal['夏普比率']):.3f}），"
            f"再平衡 {int(best_rebal['再平衡次数'])} 次，平均换手率 {float(best_rebal['平均换手率%']):.1f}%。"
        )

    # 状态转换
    if bundle.regime_df is not None and len(bundle.regime_df):
        hs300_row = bundle.regime_df[bundle.regime_df["资产"] == "沪深300"]
        if len(hs300_row):
            r = hs300_row.iloc[0]
            lines.append(
                f"市场状态：A股（沪深300）当前处于'{r['当前状态']}'，"
                f"Hurst={r['Hurst指数']}，建议策略：{r['策略建议']}。"
            )

    return lines if lines else ["量化模型数据暂未生成，请先运行各模型脚本。"]


def _make_quant_summary_table(bundle: ReportBundle) -> pd.DataFrame:
    """六模型快照汇总表"""
    rows = []

    if bundle.dcc_latest is not None:
        avg_corr = _safe_float(bundle.dcc_latest.get("平均相关系数"))
        rows.append({
            "模型": "DCC 动态相关",
            "关键指标": f"平均相关系数 {avg_corr:.3f}" if avg_corr else "-",
            "状态/结论": str(bundle.dcc_latest.get("预警状态", "-")),
        })

    if bundle.cta_df is not None and len(bundle.cta_df):
        strongest = bundle.cta_df.sort_values("合成信号", key=abs, ascending=False).iloc[0]
        rows.append({
            "模型": "CTA 趋势跟踪",
            "关键指标": f"最强: {strongest['资产']} {float(strongest['合成信号']):+.3f}",
            "状态/结论": str(strongest["操作建议"]),
        })

    if bundle.regime_df is not None and len(bundle.regime_df):
        hs300_row = bundle.regime_df[bundle.regime_df["资产"] == "沪深300"]
        if len(hs300_row):
            r = hs300_row.iloc[0]
            rows.append({
                "模型": "市场状态转换",
                "关键指标": f"Hurst={r['Hurst指数']}  AC={r['自相关(1日)']}",
                "状态/结论": f"{r['当前状态']} → {r['策略建议']}",
            })

    if bundle.rebalance_df is not None and len(bundle.rebalance_df):
        best = bundle.rebalance_df.sort_values("夏普比率", ascending=False).iloc[0]
        rows.append({
            "模型": "再平衡策略",
            "关键指标": f"最优: {best['策略']}  夏普 {float(best['夏普比率']):.3f}",
            "状态/结论": f"再平衡 {int(best['再平衡次数'])} 次，换手 {float(best['平均换手率%']):.1f}%",
        })

    if bundle.risk_parity_df is not None and len(bundle.risk_parity_df):
        top = bundle.risk_parity_df.sort_values("风险平价权重%", ascending=False).iloc[0]
        rows.append({
            "模型": "风险平价/预算",
            "关键指标": f"最大权重: {top['资产']} {float(top['风险平价权重%']):.1f}%",
            "状态/结论": "等风险贡献，动态象限调整",
        })

    if bundle.performance_df is not None and len(bundle.performance_df):
        perf_col = "资产/组合" if "资产/组合" in bundle.performance_df.columns else "资产"
        asset_perf = bundle.performance_df[~bundle.performance_df[perf_col].str.contains("组合", na=False)]
        if len(asset_perf):
            best = asset_perf.sort_values("夏普比率", ascending=False).iloc[0]
            rows.append({
                "模型": "绩效评估",
                "关键指标": f"最优: {best[perf_col]}  夏普 {float(best['夏普比率']):.3f}",
                "状态/结论": f"年化收益 {float(best['年化收益%']):.2f}%",
            })

    return pd.DataFrame(rows) if rows else pd.DataFrame(columns=["模型", "关键指标", "状态/结论"])


def _strategy_recommendations(bundle: ReportBundle) -> List[str]:
    bond_pref = _safe_float(bundle.merrill_latest["债券偏好"]) or 0.0

    # 债券偏好 [-1, +1]：正数=偏好债券=拉长久期，负数=不偏好债券=缩短久期
    if bond_pref > 0.3:
        duration = "偏长"
        duration_range = "7Y-10Y 甚至 30Y 超长端"
        duration_reason = "美林时钟显示债券偏好强（+{:.2f}），增长放缓+通胀回落环境利好长久期".format(bond_pref)
    elif bond_pref > 0:
        duration = "中性偏长"
        duration_range = "5Y-10Y"
        duration_reason = "美林时钟显示债券偏好温和（+{:.2f}），可适度拉长久期".format(bond_pref)
    elif bond_pref > -0.3:
        duration = "中性偏短"
        duration_range = "3Y-5Y"
        duration_reason = "美林时钟显示债券偏好偏弱（{:.2f}），建议缩短久期控制利率风险".format(bond_pref)
    else:
        duration = "偏短"
        duration_range = "1Y-3Y 短端"
        duration_reason = "美林时钟显示债券偏好很弱（{:.2f}），通胀上行+增长复苏不利于债券，大幅缩短久期".format(bond_pref)

    # 信用策略：根据 AA-AAA 等级利差分位数动态调整
    aa_aaa_pct = bundle.percentiles.get("aa_aaa_1y") or 50.0
    if aa_aaa_pct > 70:
        credit_rec = "等级利差偏高，下沉信用风险补偿充足但流动性折价大，建议以高等级城投、央国企产业债为主，AA 及以下严格个券筛选。"
    elif aa_aaa_pct < 30:
        credit_rec = "等级利差已明显压缩，下沉信用性价比下降，建议维持高等级配置，不宜追逐利差压缩。"
    else:
        credit_rec = "等级利差处于中性区间，可适度参与高等级城投和央国企产业债，流动性较好的二永债可作为补充。"

    # 交易策略：根据 10Y 分位数和资金面动态调整
    y10_pct = bundle.percentiles.get("y10_1y") or 50.0
    dr007 = float(bundle.rate_latest["dr007"])
    repo = float(bundle.rate_latest["reverse_repo_7d"])
    if y10_pct < 25:
        trade_rec = "10Y 利率处于历史低位，追多性价比低，建议以持有为主，等待回调再加仓，保留充足的回撤加仓空间。"
    elif y10_pct > 75:
        trade_rec = "10Y 利率处于历史高位，配置价值突出，可积极布局，适度拉长久期。"
    else:
        if dr007 < repo:
            trade_rec = "利率处于中性区间，资金面偏松，短端有支撑，可维持中性仓位，等待方向明确后再加减仓。"
        else:
            trade_rec = "利率处于中性区间，资金面偏紧，短端有压力，建议维持中性偏保守仓位，控制久期风险。"

    # 风险跟踪：根据 Crisis Score 分项动态生成
    fx_z     = _safe_float(bundle.crisis_latest.get("汇率波动率z"))  or 0.0
    commod_z = _safe_float(bundle.crisis_latest.get("商品波动率z"))  or 0.0
    equity_z = _safe_float(bundle.crisis_latest.get("股市波动率z"))  or 0.0
    risk_items = []
    if abs(fx_z) > 1.5:
        risk_items.append(f"人民币汇率波动（z={fx_z:.2f}）")
    if abs(commod_z) > 1.5:
        risk_items.append(f"商品价格剧烈波动（z={commod_z:.2f}）")
    if abs(equity_z) > 1.5:
        risk_items.append(f"股市波动率抬升（z={equity_z:.2f}）")
    if risk_items:
        risk_rec = f"重点跟踪：{'、'.join(risk_items)}，关注是否向国内债市风险偏好传导。"
    else:
        risk_rec = "当前各风险分项均处于正常区间，维持常规风险监控即可，重点关注月度宏观数据变化。"

    return [
        f"久期策略：维持{duration}，重点在 {duration_range} 做配置。{duration_reason}。",
        f"信用策略：{credit_rec}",
        f"交易策略：{trade_rec}",
        f"风险跟踪：{risk_rec}",
    ]


def _make_rate_table(bundle: ReportBundle) -> pd.DataFrame:
    rows = []
    for label, col in [
        ("1Y国债", "treasury_1y"),
        ("3Y国债", "treasury_3y"),
        ("5Y国债", "treasury_5y"),
        ("10Y国债", "treasury_10y"),
        ("30Y国债", "treasury_30y"),
    ]:
        rows.append(
            {
                "期限": label,
                "最新(%)": _fmt_num(float(bundle.rate_latest[col]), 4),
                "1个月前(%)": _fmt_num(float(bundle.rate_1m[col]), 4),
                "1个月变化(bp)": _fmt_change_bp_pct(float(bundle.rate_latest[col]), float(bundle.rate_1m[col])),
                "3个月前(%)": _fmt_num(float(bundle.rate_3m[col]), 4),
                "3个月变化(bp)": _fmt_change_bp_pct(float(bundle.rate_latest[col]), float(bundle.rate_3m[col])),
            }
        )
    rows.append(
        {
            "期限": "10Y-1Y利差",
            "最新(%)": _fmt_num(float(bundle.rate_latest["term_spread_10y_1y"]), 2),
            "1个月前(%)": _fmt_num(float(bundle.rate_1m["term_spread_10y_1y"]), 2),
            "1个月变化(bp)": _fmt_change_bp_level(float(bundle.rate_latest["term_spread_10y_1y"]), float(bundle.rate_1m["term_spread_10y_1y"])),
            "3个月前(%)": _fmt_num(float(bundle.rate_3m["term_spread_10y_1y"]), 2),
            "3个月变化(bp)": _fmt_change_bp_level(float(bundle.rate_latest["term_spread_10y_1y"]), float(bundle.rate_3m["term_spread_10y_1y"])),
        }
    )
    return pd.DataFrame(rows)


def _make_credit_table(bundle: ReportBundle) -> pd.DataFrame:
    rows = []
    for label, col in [
        ("AAA 3Y信用利差", "credit_spread_aaa_3y"),
        ("AA 3Y信用利差", "credit_spread_aa_3y"),
        ("AA-AAA等级利差", "aa_aaa_spread_3y"),
        ("国开-国债10Y利差", "cdb_treasury_spread_10y"),
    ]:
        latest = bundle.credit_latest.get(col)
        prev_1m = bundle.credit_1m.get(col)
        prev_3m = bundle.credit_3m.get(col)
        latest_val = None if pd.isna(latest) else float(latest)
        prev_1m_val = None if pd.isna(prev_1m) else float(prev_1m)
        prev_3m_val = None if pd.isna(prev_3m) else float(prev_3m)
        rows.append(
            {
                "指标": label,
                "最新(bp)": _fmt_num(latest_val, 2),
                "1个月前(bp)": _fmt_num(prev_1m_val, 2),
                "1个月变化(bp)": _fmt_change_bp_level(latest_val, prev_1m_val),
                "3个月前(bp)": _fmt_num(prev_3m_val, 2),
                "3个月变化(bp)": _fmt_change_bp_level(latest_val, prev_3m_val),
            }
        )
    return pd.DataFrame(rows)


def _make_macro_table(bundle: ReportBundle) -> pd.DataFrame:
    latest = bundle.macro_latest
    prev = bundle.macro_prev
    rows = []
    for label, col in [("PMI", "pmi"), ("CPI同比", "cpi_yoy"), ("PPI同比", "ppi_yoy"), ("M2同比", "m2_yoy"), ("社融存量同比", "social_financing_yoy")]:
        latest_val = float(latest[col])
        prev_val = float(prev[col])
        rows.append({"指标": label, "最新值": _fmt_num(latest_val, 1), "前值": _fmt_num(prev_val, 1), "变化": _fmt_change_bp_level(latest_val, prev_val), "最新日期": latest["trade_date"].strftime("%Y-%m-%d")})
    return pd.DataFrame(rows)


def _make_toolkit_table(bundle: ReportBundle) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "模型": "中国版美林时钟",
                "最新日期": str(bundle.merrill_latest["日期"]),
                "关键信号": f"增长 {bundle.merrill_latest['增长动量']} / 通胀 {bundle.merrill_latest['通胀动量']} / 流动性 {bundle.merrill_latest['流动性动量']}",
                "结论": f"{bundle.merrill_latest['传统象限']}，债券偏好 {bundle.merrill_latest['债券偏好']}",
            },
            {
                "模型": "中国版 Crisis Score",
                "最新日期": str(bundle.crisis_latest["日期"]),
                "关键信号": f"Score={bundle.crisis_latest['Crisis Score']}，汇率z={bundle.crisis_latest['汇率波动率z']}，商品z={bundle.crisis_latest['商品波动率z']}",
                "结论": f"{bundle.crisis_latest['市场状态']}，{bundle.crisis_latest['操作建议']}",
            },
        ]
    )


def _make_garch_table(bundle: ReportBundle) -> pd.DataFrame:
    rows = []
    for _, row in bundle.garch_df.iterrows():
        rows.append(
            {
                "资产": row["asset_name"],
                "最优模型": row["最优模型"],
                "实际α": f'{row['alpha']:.3f}',
                "默认α": f'{row['note_alpha']:.2f}',
                "实际β": f'{row['beta']:.3f}',
                "默认β": f'{row['note_beta']:.2f}',
                "年化波动率%": f'{row['年化波动率%']:.2f}',
                "样本外相关性": f'{row['样本外相关性']:.3f}',
                "状态": row["波动率状态"],
            }
        )
    return pd.DataFrame(rows)


def _make_cross_model_table(bundle: ReportBundle) -> pd.DataFrame:
    garch_extreme = bundle.garch_df[bundle.garch_df["波动率状态"].astype(str).str.contains("极端")]["asset_name"].tolist()
    garch_signal = f"{'、'.join(garch_extreme)} 极端波动" if garch_extreme else "无极端波动资产"
    garch_meaning = "高波动资产需控制仓位，风险预算优先防御" if garch_extreme else "波动率正常，可按计划配置"

    crisis_score = _safe_float(bundle.crisis_latest.get("Crisis Score")) or 0.0
    fx_z     = _safe_float(bundle.crisis_latest.get("汇率波动率z"))  or 0.0
    if crisis_score < 1:
        crisis_meaning = "系统性风险可控，无需大幅防御"
    elif crisis_score < 2:
        crisis_meaning = "风险在积累，建议适度降低风险敞口"
    else:
        crisis_meaning = "高风险区间，建议大幅防御"

    merrill_phase = str(bundle.merrill_latest["传统象限"])
    prefs = [
        ("黄金", _safe_float(bundle.merrill_latest["黄金偏好"]) or 0.0),
        ("股票", _safe_float(bundle.merrill_latest["股票偏好"]) or 0.0),
        ("债券", _safe_float(bundle.merrill_latest["债券偏好"]) or 0.0),
        ("现金", _safe_float(bundle.merrill_latest["现金偏好"]) or 0.0),
        ("商品", _safe_float(bundle.merrill_latest["商品偏好"]) or 0.0),
    ]
    top3 = " > ".join(n for n, _ in sorted(prefs, key=lambda x: x[1], reverse=True)[:3])
    phase_meanings = {
        "复苏": "增长回升通胀低，利好股票和信用债",
        "过热": "增长通胀双升，利好商品，债券承压",
        "滞胀": "增长弱通胀升，利好黄金现金，债券股票均承压",
        "衰退": "增长通胀双降，利好债券，股票商品承压",
    }
    merrill_meaning = phase_meanings.get(merrill_phase, f"{merrill_phase}象限")

    return pd.DataFrame([
        {"模型": "GARCH", "当前信号": garch_signal, "含义": garch_meaning},
        {"模型": "Crisis Score", "当前信号": f"{crisis_score:.3f}，汇率z={fx_z:.2f}", "含义": crisis_meaning},
        {"模型": "美林时钟", "当前信号": f"{merrill_phase}，{top3}", "含义": merrill_meaning},
    ])


def _build_document(bundle: ReportBundle, chart_paths: List[Path]) -> Document:
    document = Document()
    _set_document_style(document)
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    subtitle = (
        f"生成日期：{datetime.now().strftime('%Y-%m-%d')} | 债券市场最新数据截至 {bundle.rate_latest['trade_date'].strftime('%Y-%m-%d')} "
        f"| 月度宏观最新数据截至 {bundle.macro_latest['trade_date'].strftime('%Y-%m-%d')}"
    )
    _add_title(document, "债券及宏观观察报告", subtitle)
    document.add_paragraph("数据来源：MOSS 本地 market.db（日频利率、信用、流动性与月度宏观） + macro_toolkit 输出（GARCH、Crisis Score、美林时钟）。")

    # 核心结论：全部动态生成
    pmi = float(bundle.macro_latest["pmi"])
    cpi = float(bundle.macro_latest["cpi_yoy"])
    ppi = float(bundle.macro_latest["ppi_yoy"])
    y10 = float(bundle.rate_latest["treasury_10y"])
    aa_aaa = float(bundle.credit_latest["aa_aaa_spread_3y"])
    aa_aaa_pct = bundle.percentiles.get("aa_aaa_1y") or 50.0
    crisis_score = _safe_float(bundle.crisis_latest.get("Crisis Score")) or 0.0
    merrill_phase = str(bundle.merrill_latest["传统象限"])

    bond_pref = _safe_float(bundle.merrill_latest["债券偏好"]) or 0.0
    if bond_pref > 0.3:
        bond_view = f"美林时钟债券偏好 {bond_pref:+.2f}，长端配置价值突出，可积极拉长久期"
    elif bond_pref > 0:
        bond_view = f"美林时钟债券偏好 {bond_pref:+.2f}，长端有配置价值，可适度拉长久期"
    elif bond_pref > -0.3:
        bond_view = f"美林时钟债券偏好 {bond_pref:.2f}，偏弱，建议缩短久期至 3Y-5Y，控制利率风险"
    else:
        bond_view = f"美林时钟债券偏好 {bond_pref:.2f}，明显偏弱，建议大幅缩短久期至短端，规避利率上行风险"

    # 增长描述
    growth_view = "增长端偏弱" if pmi < 50 else "增长端温和扩张"
    # 价格描述
    if cpi < 0 and ppi < 0:
        price_view = "通缩压力较重，价格环境利好债券"
    elif cpi < 1 and ppi < 0:
        price_view = "价格低位，通缩压力尚未消散，对债券无明显逆风"
    elif cpi >= 2 or ppi >= 3:
        price_view = "价格动量明显抬升，通胀压力对债券形成压制"
    else:
        price_view = "价格温和，通胀压力可控，对债券中性"
    # 信用描述
    if aa_aaa_pct > 70:
        credit_view = f"AA-AAA 等级利差 {aa_aaa:.2f}bp（近1年{aa_aaa_pct:.0f}%分位），下沉信用风险补偿偏高，需控制节奏"
    elif aa_aaa_pct < 30:
        credit_view = f"AA-AAA 等级利差 {aa_aaa:.2f}bp（近1年{aa_aaa_pct:.0f}%分位），等级利差已压缩，不宜继续下沉"
    else:
        credit_view = f"AA-AAA 等级利差 {aa_aaa:.2f}bp（近1年{aa_aaa_pct:.0f}%分位），信用分层合理"
    # 模型综合
    extreme = bundle.garch_df[bundle.garch_df["波动率状态"].astype(str).str.contains("极端")]["asset_name"].tolist()
    if crisis_score >= 2:
        model_view = f"三模型高度防御：GARCH 极端波动资产{'、'.join(extreme) if extreme else '无'}，Crisis Score {crisis_score:.3f}（高风险），美林时钟{merrill_phase}"
    elif crisis_score >= 1 or extreme:
        model_view = f"三模型偏防御：{'GARCH 极端波动资产' + '、'.join(extreme) + '，' if extreme else ''}Crisis Score {crisis_score:.3f}（警惕），美林时钟{merrill_phase}"
    else:
        model_view = f"三模型信号中性：Crisis Score {crisis_score:.3f}（正常），美林时钟{merrill_phase}，无极端波动预警"

    document.add_heading("一、核心结论", level=1)
    _add_bullets(
        document,
        [
            f"{growth_view}，PMI {pmi:.1f}；{price_view}，CPI {cpi:.1f}%，PPI {ppi:.1f}%。",
            f"截至 {bundle.rate_latest['trade_date'].strftime('%Y-%m-%d')}，10Y 国债 {y10:.4f}%。{bond_view}。",
            f"信用层面：{credit_view}。",
            f"{model_view}。",
        ],
    )
    document.add_heading("二、宏观环境", level=1)
    document.add_paragraph("月度宏观指标表")
    _add_table(document, _make_macro_table(bundle))
    document.add_paragraph("")
    for text in _macro_takeaways(bundle):
        document.add_paragraph(text)
    document.add_picture(str(chart_paths[3]), width=Inches(6.8))
    document.add_heading("三、利率债市场", level=1)
    document.add_paragraph("国债关键期限利率表")
    _add_table(document, _make_rate_table(bundle))
    document.add_paragraph("")
    for text in _bond_takeaways(bundle):
        document.add_paragraph(text)
    document.add_picture(str(chart_paths[0]), width=Inches(6.8))
    document.add_picture(str(chart_paths[1]), width=Inches(6.8))
    document.add_heading("四、信用债与风险偏好", level=1)
    document.add_paragraph("信用利差表")
    _add_table(document, _make_credit_table(bundle))
    document.add_paragraph("")
    for text in _credit_takeaways(bundle):
        document.add_paragraph(text)
    document.add_picture(str(chart_paths[2]), width=Inches(6.8))
    document.add_heading("五、量化模型增补", level=1)
    document.add_paragraph("模型摘要")
    _add_table(document, _make_toolkit_table(bundle))
    document.add_paragraph("")
    document.add_paragraph("GARCH 参数对照")
    _add_table(document, _make_garch_table(bundle))
    document.add_paragraph("")
    for text in _garch_takeaways(bundle):
        document.add_paragraph(text)
    document.add_picture(str(chart_paths[4]), width=Inches(6.8))
    document.add_picture(str(chart_paths[5]), width=Inches(6.8))
    document.add_heading("六、模型交叉验证", level=1)
    _add_table(document, _make_cross_model_table(bundle))
    document.add_paragraph("")
    for text in _model_cross_takeaways(bundle):
        document.add_paragraph(text)
    document.add_heading("七、量化模型扩展", level=1)
    document.add_paragraph("六大量化模型快照")
    _add_table(document, _make_quant_summary_table(bundle))
    document.add_paragraph("")
    for text in _quant_models_takeaways(bundle):
        document.add_paragraph(text)
    # 插入已生成的图表（如果存在）
    for img_name in [
        "dcc_heatmap.png",
        "dcc_timeseries.png",
        "cta_signals.png",
        "regime_overview.png",
        "regime_distribution.png",
        "rebalance_comparison.png",
        "weight_drift.png",
    ]:
        img_path = ASSET_DIR / img_name
        if img_path.exists():
            document.add_picture(str(img_path), width=Inches(6.8))
    document.add_heading("八、全策略综合回测", level=1)
    if bundle.backtest_df is not None and len(bundle.backtest_df):
        document.add_paragraph("策略绩效汇总（近5年，扣除交易成本）")
        _add_table(document, bundle.backtest_df)
        document.add_paragraph("")
        if bundle.backtest_annual_df is not None:
            document.add_paragraph("分年度收益 (%)")
            _add_table(document, bundle.backtest_annual_df)
            document.add_paragraph("")
        # 回测图表
        for img_name in ["backtest_nav.png", "backtest_annual.png", "backtest_metrics.png"]:
            img_path = ASSET_DIR / img_name
            if img_path.exists():
                document.add_picture(str(img_path), width=Inches(6.8))
        # 回测结论
        best_row = bundle.backtest_df.sort_values("夏普比率", ascending=False).iloc[0]
        document.add_paragraph(
            f"回测结论：近5年最优策略为'{best_row['策略']}'，夏普比率 {float(best_row['夏普比率']):.3f}，"
            f"年化收益 {float(best_row['年化收益%']):.2f}%，最大回撤 {float(best_row['最大回撤%']):.2f}%。"
            f"全模型综合策略通过风险平价基础权重 + CTA信号微调 + 状态防御，在控制回撤的同时保持了较好的风险调整收益。"
        )
    else:
        document.add_paragraph("回测数据暂未生成，请先运行 backtest_cn.py。")
    document.add_heading("九、策略建议", level=1)
    _add_bullets(document, _strategy_recommendations(bundle))
    document.add_heading("十、说明", level=1)
    document.add_paragraph("1. 债券日频数据采用 market.db 中最新有效交易日。 2. 月度宏观指标采用数据库内最新非空发布值，日期可能晚于实际统计月份。 3. 本报告用于内部分析，不构成投资建议。")
    document.add_section(WD_SECTION.NEW_PAGE)
    return document


def main() -> None:
    _set_matplotlib_style()
    bundle = _build_bundle()
    chart_paths = [
        _plot_yield_curve(bundle),
        _plot_liquidity(bundle),
        _plot_credit(bundle),
        _plot_macro(bundle),
        _plot_garch_params(bundle),
        _plot_garch_risk(bundle),
    ]
    document = _build_document(bundle, chart_paths)
    document.save(OUTPUT_DOC)
    print(f"Word report saved to: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
